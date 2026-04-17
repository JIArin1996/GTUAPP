import re
import tempfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from flask import Flask, after_this_request, jsonify, render_template, request, send_file
from flask_sqlalchemy import SQLAlchemy
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF

from parser_utils import procesar_entrada, sanitizar_nombre_archivo

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///gtuapp.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

# --- Modelos ---

class Cliente(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    razon_social = db.Column(db.String(200), nullable=False)
    domicilio_constituido = db.Column(db.Text, nullable=False)
    nombre_cliente = db.Column(db.String(200))
    ci = db.Column(db.String(50))
    rut = db.Column(db.String(50))
    telefono = db.Column(db.String(50))
    mail = db.Column(db.String(100))
    fideicomiso = db.Column(db.String(200)) # NUEVO CAMPO
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)

    def to_dict(self):
        return {
            "id": self.id,
            "razon_social": self.razon_social,
            "domicilio_constituido": self.domicilio_constituido,
            "nombre_cliente": self.nombre_cliente,
            "ci": self.ci,
            "rut": self.rut,
            "telefono": self.telefono,
            "mail": self.mail,
            "fideicomiso": self.fideicomiso
        }

# --- Inicialización ---

OUTPUT_ANEXOS = Path("output_anexos")
OUTPUT_ANEXOS.mkdir(exist_ok=True)

with app.app_context():
    db.create_all()
    # Migracion simple para agregar columna fideicomiso si no existe
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text("ALTER TABLE cliente ADD COLUMN fideicomiso VARCHAR(200)"))
            conn.commit()
    except Exception:
        # Si ya existe, fallara silenciosamente
        pass

# --- Constantes y Helpers ---
SNIG_PATTERN = re.compile(r"(?<!\d)(8580000\d{8})(?!\d)")
TXT_SUFFIX = "|.|.|.|.|.|.|.|.|.|.|]"
TXT_PREFIX = "A0000000"

MESES_ES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto", 
    9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

def format_date_spanish(date_str, format_type="long"):
    """
    Formatea 'YYYY-MM-DD' a:
    'long': '17 de Abril del 2026'
    'filename': '17abr2026'
    """
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        day = dt.day
        month_name = MESES_ES[dt.month]
        year = dt.year
        
        if format_type == "filename":
            short_month = month_name[:3].lower()
            return f"{day:02d}{short_month}{year}"
        
        return f"{day} de {month_name} del {year}"
    except Exception:
        return date_str


def leer_pdf(ruta_pdf: Path) -> str:
    with ruta_pdf.open("rb") as archivo:
        lector = PdfReader(archivo)
        texto = "\n".join([pagina.extract_text() or "" for pagina in lector.pages])
    return texto.strip()


def generar_excel(datos, ruta_excel: Path) -> Path:
    df = pd.DataFrame(datos, columns=["Caravana", "Sexo", "Edad"])
    df.to_excel(ruta_excel, index=False)
    return ruta_excel


def es_peticion_ajax() -> bool:
    return request.headers.get("X-Requested-With") == "XMLHttpRequest"


def respuesta_error(mensaje: str, status_code: int = 400):
    if es_peticion_ajax():
        return jsonify({"ok": False, "error": mensaje}), status_code
    return render_template("index.html", error=mensaje), status_code


def extraer_caravanas_snig_excel(ruta_excel: Path):
    hojas = pd.read_excel(
        ruta_excel,
        sheet_name=None,
        header=None,
        dtype=str,
        engine="openpyxl",
    )

    caravanas = []
    vistas = set()

    for _, df in hojas.items():
        if df is None or df.empty:
            continue

        for fila in df.itertuples(index=False, name=None):
            for celda in fila:
                if celda is None:
                    continue
                for match in SNIG_PATTERN.findall(str(celda)):
                    if match not in vistas:
                        vistas.add(match)
                        caravanas.append(match)

    return caravanas


def construir_txt_snig(caravanas, guia: str, ahora: datetime | None = None) -> str:
    if ahora is None:
        ahora = datetime.now()
    fecha = ahora.strftime("%d%m%Y")
    hora = ahora.strftime("%H%M")
    lineas = [
        f"[|{TXT_PREFIX}{caravana}|{fecha}|{hora}|{guia}{TXT_SUFFIX}"
        for caravana in caravanas
    ]
    return "\n".join(lineas) + "\n"


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        nombre_archivo = sanitizar_nombre_archivo(request.form.get("nombre_archivo", ""))
        archivos_pdf = request.files.getlist("archivo_pdf")

        if not nombre_archivo:
            return respuesta_error("Ingresá un nombre de archivo válido.")

        archivos_validos = [archivo for archivo in archivos_pdf if archivo and archivo.filename]
        if not archivos_validos:
            return respuesta_error("Debes cargar al menos un archivo PDF.")

        if any(not archivo.filename.lower().endswith(".pdf") for archivo in archivos_validos):
            return respuesta_error("Todos los archivos cargados deben ser PDF.")

        try:
            temp_dir = Path(tempfile.gettempdir())
            rutas_pdf = []
            datos_procesados = []

            for idx, archivo_pdf in enumerate(archivos_validos):
                ruta_pdf = temp_dir / f"pdf_temporal_gtu_{idx}.pdf"
                archivo_pdf.save(ruta_pdf)
                rutas_pdf.append(ruta_pdf)

                texto_entrada = leer_pdf(ruta_pdf)
                datos_procesados.extend(procesar_entrada(texto_entrada))

            if not datos_procesados:
                return respuesta_error("No se encontraron datos con el formato esperado en los PDF.")

            tmp_excel = tempfile.NamedTemporaryFile(prefix="gtu_", suffix=".xlsx", delete=False)
            ruta_excel = Path(tmp_excel.name)
            tmp_excel.close()

            generar_excel(datos_procesados, ruta_excel)

            @after_this_request
            def cleanup_temporales(response):
                for ruta_pdf in rutas_pdf:
                    try:
                        if ruta_pdf.exists():
                            ruta_pdf.unlink()
                    except OSError:
                        pass

                try:
                    if ruta_excel.exists():
                        ruta_excel.unlink()
                except OSError:
                    pass

                return response

            return send_file(
                ruta_excel,
                as_attachment=True,
                download_name=f"{nombre_archivo}.xlsx",
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception as exc:
            return respuesta_error(f"Ocurrió un error al procesar el archivo: {exc}", status_code=500)

    return render_template("index.html")


@app.route("/excel-a-txt", methods=["POST"])
def excel_a_txt():
    archivo_excel = request.files.get("archivo_excel")
    guia = (request.form.get("guia") or "").strip()
    nombre_archivo = sanitizar_nombre_archivo(request.form.get("nombre_txt", ""))

    if not archivo_excel or not archivo_excel.filename:
        return respuesta_error("Debes cargar un archivo Excel (.xlsx).")

    if not archivo_excel.filename.lower().endswith(".xlsx"):
        return respuesta_error("El archivo debe tener extensión .xlsx.")

    if not guia:
        return respuesta_error("Debes ingresar el número de guía.")

    if not re.fullmatch(r"[A-Za-z]\d{6}", guia):
        return respuesta_error("El número de guía debe tener 1 letra seguida de 6 números (ej: D674195).")

    if not nombre_archivo:
        nombre_archivo = "salida_snig"

    ruta_excel = None
    ruta_txt = None

    try:
        tmp_excel = tempfile.NamedTemporaryFile(prefix="excel_temporal_gtu_", suffix=".xlsx", delete=False)
        ruta_excel = Path(tmp_excel.name)
        tmp_excel.close()
        archivo_excel.save(ruta_excel)

        caravanas = extraer_caravanas_snig_excel(ruta_excel)
        if not caravanas:
            return respuesta_error(
                "No se encontraron caravanas SNIG válidas (15 dígitos que comiencen con 8580000)."
            )

        contenido_txt = construir_txt_snig(caravanas, guia)

        tmp_txt = tempfile.NamedTemporaryFile(prefix="gtu_snig_", suffix=".txt", delete=False)
        ruta_txt = Path(tmp_txt.name)
        tmp_txt.close()
        ruta_txt.write_text(contenido_txt, encoding="utf-8")

        @after_this_request
        def cleanup_temporales_excel_txt(response):
            try:
                if ruta_excel and ruta_excel.exists():
                    ruta_excel.unlink()
            except OSError:
                pass

            try:
                if ruta_txt and ruta_txt.exists():
                    ruta_txt.unlink()
            except OSError:
                pass

            return response

        download_name = f"{nombre_archivo.replace('.txt', '')}.txt"
        return send_file(
            ruta_txt,
            as_attachment=True,
            download_name=download_name,
            mimetype="text/plain; charset=utf-8",
        )
    except Exception as exc:
        return respuesta_error(f"Ocurrió un error al procesar el Excel: {exc}", status_code=500)


@app.route("/txt-a-excel", methods=["POST"])
def txt_a_excel():
    archivos_txt = request.files.getlist("archivo_txt")
    nombre_archivo = sanitizar_nombre_archivo(request.form.get("nombre_excel", ""))

    if not nombre_archivo:
        nombre_archivo = "caravanas_extraidas"

    archivos_validos = [archivo for archivo in archivos_txt if archivo and archivo.filename]
    if not archivos_validos:
        return respuesta_error("Debes cargar al menos un archivo TXT.")

    if any(not archivo.filename.lower().endswith(".txt") for archivo in archivos_validos):
        return respuesta_error("Todos los archivos cargados deben ser TXT.")

    try:
        from appgtu_extra import extraer_caravanas_txt # Si moviste esta funcion, ajustala
        # Como no esta en la version actual de appgtu.py en esta rama, la re-implemento aqui si es necesario
        # o la tomo de la version anterior si la tengo.
        
        # Implementacion directa para asegurar que funcione:
        def local_extraer_caravanas_txt(archivos):
            res = []
            ptrn = re.compile(r"^.*?\[\|A0000000(8580000\d{8})")
            for f in archivos:
                lineas = f.read().decode("utf-8-sig", errors="ignore").splitlines()
                fname = f.filename
                for l in lineas:
                    m = ptrn.search(l.strip())
                    if m: res.append((m.group(1), fname))
            return res

        datos_extraidos = local_extraer_caravanas_txt(archivos_validos)

        if not datos_extraidos:
            return respuesta_error("No se encontraron caravanas válidas en los archivos TXT.")

        tmp_excel = tempfile.NamedTemporaryFile(prefix="gtu_txt_", suffix=".xlsx", delete=False)
        ruta_excel = Path(tmp_excel.name)
        tmp_excel.close()

        df = pd.DataFrame(datos_extraidos, columns=["Caravana", "Archivo Origen"])
        df.to_excel(ruta_excel, index=False)

        @after_this_request
        def cleanup_temporales_txt_excel(response):
            try:
                if ruta_excel.exists():
                    ruta_excel.unlink()
            except OSError:
                pass
            return response

        return send_file(
            ruta_excel,
            as_attachment=True,
            download_name=f"{nombre_archivo}.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except Exception as exc:
        return respuesta_error(f"Ocurrió un error al procesar los TXT: {exc}", status_code=500)


@app.route("/anexos/clientes", methods=["GET", "POST"])
def gestionar_clientes():
    if request.method == "POST":
        data = request.json
        if not data.get("razon_social") or not data.get("domicilio_constituido"):
            return jsonify({"ok": False, "error": "Razón Social y Domicilio son obligatorios"}), 400
        
        nuevo = Cliente(
            razon_social=data["razon_social"],
            domicilio_constituido=data["domicilio_constituido"],
            nombre_cliente=data.get("nombre_cliente"),
            ci=data.get("ci"),
            rut=data.get("rut"),
            telefono=data.get("telefono"),
            mail=data.get("mail"),
            fideicomiso=data.get("fideicomiso") # GUARDAR FIDEICOMISO
        )
        db.session.add(nuevo)
        db.session.commit()
        return jsonify({"ok": True, "cliente": nuevo.to_dict()})
    
    # Filtrar por fideicomiso si se proporciona
    fid_filter = request.args.get("fideicomiso")
    query = Cliente.query
    if fid_filter:
        query = query.filter_by(fideicomiso=fid_filter)
        
    clientes = query.order_by(Cliente.razon_social).all()
    return jsonify([c.to_dict() for c in clientes])


@app.route("/anexos/clientes/<int:id>", methods=["PUT"])
def editar_cliente(id):
    cliente = Cliente.query.get_or_404(id)
    data = request.json
    cliente.razon_social = data.get("razon_social", cliente.razon_social)
    cliente.domicilio_constituido = data.get("domicilio_constituido", cliente.domicilio_constituido)
    cliente.nombre_cliente = data.get("nombre_cliente", cliente.nombre_cliente)
    cliente.ci = data.get("ci", cliente.ci)
    cliente.rut = data.get("rut", cliente.rut)
    cliente.telefono = data.get("telefono", cliente.telefono)
    cliente.mail = data.get("mail", cliente.mail)
    db.session.commit()
    return jsonify({"ok": True, "cliente": cliente.to_dict()})


def generar_documento(datos, custom_filename=None):
    doc = Document()
    doc.add_heading(f"ANEXO - {datos['fideicomiso']}", 0)
    
    p = doc.add_paragraph("Este es un documento generado automáticamente como placeholder.\n\n")
    p.add_run(f"Tipo de Anexo: {datos['tipo_anexo']}\n")
    p.add_run(f"Cliente: {datos['cliente']['razon_social']}\n")
    p.add_run(f"Domicilio: {datos['cliente']['domicilio_constituido']}\n")
    p.add_run(f"Fecha: {format_date_spanish(datos['fecha'])}\n")
    
    doc.add_paragraph("\nPróximamente se integrarán las plantillas legales oficiales.")
    
    if custom_filename:
        filename = custom_filename
        if not filename.lower().endswith(".docx"):
            filename += ".docx"
    else:
        # Formato nombre: AnexoNRO_NombreCliente_DDmesAAAA
        n_cli = datos['cliente']['razon_social'].replace(" ", "_")
        f_anexo = format_date_spanish(datos['fecha'], "filename")
        n_anexo = datos['tipo_anexo'].split(" ")[1] if " " in datos['tipo_anexo'] else "X"
        filename = f"Anexo{n_anexo}_{n_cli}_{f_anexo}.docx"
    
    path = OUTPUT_ANEXOS / filename
    doc.save(path)
    return filename


def generar_pdf(datos, custom_filename=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, f"ANEXO - {datos['fideicomiso']}", ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("helvetica", "", 12)
    pdf.multi_cell(0, 10, "Este es un documento generado automáticamente como placeholder.\n\n")
    pdf.cell(0, 10, f"Tipo de Anexo: {datos['tipo_anexo']}", ln=True)
    pdf.cell(0, 10, f"Cliente: {datos['cliente']['razon_social']}", ln=True)
    pdf.cell(0, 10, f"Domicilio: {datos['cliente']['domicilio_constituido']}", ln=True)
    pdf.cell(0, 10, f"Fecha: {format_date_spanish(datos['fecha'])}", ln=True)
    
    pdf.ln(10)
    pdf.multi_cell(0, 10, "\nPróximamente se integrarán las plantillas legales oficiales.")
    
    if custom_filename:
        filename = custom_filename
        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"
    else:
        n_cli = datos['cliente']['razon_social'].replace(" ", "_")
        f_anexo = format_date_spanish(datos['fecha'], "filename")
        n_anexo = datos['tipo_anexo'].split(" ")[1] if " " in datos['tipo_anexo'] else "X"
        filename = f"Anexo{n_anexo}_{n_cli}_{f_anexo}.pdf"
    
    path = OUTPUT_ANEXOS / filename
    pdf.output(str(path))
    return filename


@app.route("/anexos/generar", methods=["POST"])
def generar_anexo():
    data = request.json
    cliente_id = data.get("cliente_id")
    if not cliente_id:
        return jsonify({"ok": False, "error": "Debe seleccionar un cliente"}), 400
    
    cliente = Cliente.query.get(cliente_id)
    if not cliente:
        return jsonify({"ok": False, "error": "Cliente no encontrado"}), 400
    
    try:
        datos_generacion = {
            "fideicomiso": data.get("fideicomiso"),
            "tipo_anexo": data.get("tipo_anexo"),
            "cliente": cliente.to_dict(),
            "fecha": data.get("fecha")
        }
        
        formato = data.get("formato", "word").lower()
        nombre_personalizado = data.get("nombre_archivo")
        
        if formato == "pdf":
            filename = generar_pdf(datos_generacion, nombre_personalizado)
        else:
            filename = generar_documento(datos_generacion, nombre_personalizado)
            
        return jsonify({"ok": True, "filename": filename})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/anexos/descargar/<filename>")
def descargar_anexo(filename):
    path = OUTPUT_ANEXOS / filename
    if not path.exists():
        return "Archivo no encontrado", 404
    return send_file(path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=False)
